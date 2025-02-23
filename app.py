import sys
import gc
# https://docs.google.com/spreadsheets/d/1Cak9V_QRjEbnlJzeEe8uj-VN864IrBpt/edit?gid=1319013326#gid=1319013326
import json
import pickle
import io
import secrets
import os
import re
import io
import base64
import random
import pytesseract
from PIL import Image
from docx import Document
from langchain.docstore import document 
from googleapiclient.errors import HttpError
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.oauth2 import service_account
from googleapiclient.discovery import build
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain_community.document_loaders import UnstructuredExcelLoader
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from langchain_community.vectorstores import FAISS
from flask import Flask,request, render_template, redirect, url_for, flash, jsonify
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.memory import ConversationSummaryBufferMemory
from functions import get_chain,get_answer,convert_pdf_to_text,sql_connection, send_mail
from exception import *
from flask import Flask, request, render_template, redirect, url_for, session, flash
pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'
app = Flask(__name__)
app.secret_key = secrets.token_hex(24)  # Required for flash messages
SERVICE_ACCOUNT_FILE = '/home/bharat/AI_Agent/ai_question_answer/GHCL_Bot/service_account.json'
# SERVICE_ACCOUNT_FILE = r'C:\Users\G01889\OneDrive\Documents\llm_project\Quest_Ans_For_All_Ext-txt-pdf.csv-1xlsx-1\midyear-pattern-444505-m8-e4bcb24550e3.json'

# credentials path
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r"C:\Users\G01889\OneDrive\Documents\llm_project\Quest_Ans_For_All_Ext-txt-pdf.csv-1xlsx-1\client_secret_820887058416-7c10c17qjh42739ca2hc0bn3cn40fdc0.apps.googleusercontent.com.json"
os.environ['SECRET_KEY'] = secrets.token_hex(5)
# Upload folder configuration
UPLOAD_FOLDER = 'Document/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['store_answer_feedback'] = 'Store_Ans'

answer_file_path = os.path.join(app.config['store_answer_feedback'], 'answers.json')
feedback_file_path = os.path.join(app.config['store_answer_feedback'], 'feedback.json')

if not os.path.exists(answer_file_path):
    with open(answer_file_path, 'w') as f:
        json.dump([], f)  # Initialize with an empty list

# Define scopes for Google Drive API
SCOPES = ['https://www.googleapis.com/auth/drive','https://www.googleapis.com/auth/spreadsheets']

# Shared credentials for all users
credentials = service_account.Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=SCOPES)

def authenticate_and_list_files():
    """Authenticate with Google Drive API and list files."""
    try:
        credentials = None
        # Check for existing token file
        if os.path.exists('token.json'):
            credentials = Credentials.from_authorized_user_file('token.json', SCOPES)
        # If no valid credentials, initiate authentication flow
        if not credentials or not credentials.valid:
            flow = InstalledAppFlow.from_client_secrets_file(
                os.environ['GOOGLE_APPLICATION_CREDENTIALS'], SCOPES)
            credentials = flow.run_local_server(port=0)
            # Save credentials for future use
            with open('token.json', 'w') as token_file:
                token_file.write(credentials.to_json())
        return credentials  # Return credentials for further use

    except Exception as e:
        return jsonify({'error':'Credential Not Found'})

def load_file_data(file_id, credentials):
    """Load data from the selected file based on its type (CSV, Excel, PDF)."""
    try:
        # Use the credentials passed into the function to authenticate the service
        service = build('drive', 'v3', credentials=credentials)
        file_metadata = service.files().get(fileId=file_id, fields="name, mimeType").execute()
        mime_type = file_metadata.get('mimeType')
        # Get the content of the file
        file_content = service.files().get_media(fileId=file_id).execute()
        
        # Check the file type
        if mime_type == 'application/vnd.ms-excel' or mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            print("Processing Excel file...")
            # Get the content of the file
            loader = UnstructuredExcelLoader(file_path=None, file=io.BytesIO(file_content))
            data = loader.load() 

        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            doc = Document(io.BytesIO(file_content))
            # Extract text from the document
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            # convert text to document
            data = [document.Document(page_content=text)]

        elif mime_type == 'application/pdf':
            text = convert_pdf_to_text(file_content)
            # convert text to document
            data = [document.Document(page_content=text)]

        elif mime_type=='text/plain' or mime_type=='text/csv':
            # convert text to document
            data = [document.Document(file_content.decode('utf-8'))]

        elif mime_type=='image/png' or mime_type=='image/jpeg':
            # Convert binary content to an image
            image = Image.open(io.BytesIO(file_content))
            text = pytesseract.image_to_string(image)
            # convert text to document
            data = [document.Document(page_content=text)]

        del service,file_metadata,mime_type,file_content
        gc.collect()
        return data
     
    except Exception as e:
        return "Not Found"
    
@app.route('/')
def index():
    return render_template('login.html')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        
        sql = f"SELECT email_id FROM public.register_table WHERE email_id = '{email}';"
        curr, connection = sql_connection()
        curr.execute(sql)
        rows = curr.fetchall()
        connection.close()

        if len(rows) == 0:
            # Check if passwords match
            if password != confirm_password:
                flash("Passwords do not match!", "error")
                return redirect(url_for('signup'))
            
            # Check password strength using regex
            password_pattern = re.compile(r'^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*?&])[A-Za-z\d@$!%*?&]{8,}$')
            if not password_pattern.match(password):
                flash("Password must contain at least 8 characters, including one uppercase letter, one lowercase letter, one digit, and one special character.", "error")
                return redirect(url_for('signup'))
            
            # Generate a random token
            token = random.randint(100000, 999999)
            # Store the token in the session for validation
            session['token'] = str(token)
            session['email'] = email
            # this required for adding pass and name after validation
            session['password'] = password
            session['name'] = name
            # Send the token via email
            subject = "Email Verification Code"
            body = f"Your verification code is {token}. Please enter it on the website to verify your email."
            message = f"Subject: {subject}\n\n{body}"
            msg = send_mail(email, message)

            if msg == 0:
                flash("Code has been sent to register email id.", "info")
                # Redirect to the validate_mail route with email as a parameter
                return redirect(url_for('validate_mail', email=email))
            del password_pattern,token,subject,body,message,msg
            gc.collect()
        else:
            flash("Email Already Exist.", "info")
    return render_template('signup.html')
@app.route('/validate_mail',methods=['POST','GET'])
def validate_mail():
    email = request.args.get('email')  # Retrieve email from query string
     
    if request.method == 'POST':
        entered_token = str(request.form['token'])

        # Compare the entered token with the session token
        if str(session.get('token')) == str(entered_token):
            password = session['password']
            name = session['name']
            encPassword = base64.b64encode(password.encode("utf-8"))
            sql = "INSERT INTO public.register_table (username, password, email_id) VALUES (%s, %s, %s);"
            curr,connection = sql_connection()
            curr.execute(sql, (name, encPassword, email))
            connection.commit()
            connection.close()

            #remove session after adding it to table 
            session.pop('password')
            session.pop('name')
            session.pop('token')

            flash("Signup successful! Please login.", "success")
            del password,name,encPassword,sql
            gc.collect()

            return redirect(url_for('login_page'))
        else:
            # return "Invalid token. Please try again.", 400
            flash("Invalid code. Please try again.", "error")  # Flash error message

    return render_template('validate_mail.html', email=email)


@app.route('/validate_mail_reset_password',methods=['POST','GET'])
def validate_mail_reset_password():
    email = request.args.get('email')  # Retrieve email from query string
    if request.method == 'POST':
        entered_token = str(request.form['token'])

        # Compare the entered token with the session token
        if str(session.get('reset_token')) == str(entered_token):
            return redirect(url_for('reset_password'))
        else:
            # return "Invalid token. Please try again.", 400
            flash("Invalid code. Please try again.", "error")  # Flash error message

    return render_template('reset_token_validate.html', email=email)


@app.route('/reset_password', methods=['GET', 'POST'])
def reset_password(): 

    if request.method == 'POST':
        email = session['email']
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']
        password_pattern = re.compile(r'^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)(?=.*[@$!%*?&])[A-Za-z\d@$!%*?&]{8,}$')
        if not password_pattern.match(new_password):
            flash("Password must contain at least 8 characters, including one uppercase letter, one lowercase letter, one digit, and one special character.", "error")
            return render_template('reset_password.html')
    
        if new_password != confirm_password:
            flash("Passwords do not match.", "error")
            return render_template('reset_password.html')
        # Check password strength using regex
        else:
            encPassword = base64.b64encode(new_password.encode("utf-8"))
            sql = "UPDATE public.register_table SET password = %s WHERE email_id = %s;"
            curr, connection = sql_connection()
            curr.execute(sql,(encPassword,email))
            connection.commit()
            connection.close()

            flash("Password has been reset successfully. You can now log in.", "success")
            del encPassword,sql
            gc.collect()
            return redirect(url_for('login_page'))
    return render_template('reset_password.html')

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email']
        session['email'] = email
        sql = f"SELECT * FROM public.register_table WHERE email_id = '{email}';"
        curr, connection = sql_connection()
        curr.execute(sql)
        rows = curr.fetchall()
        connection.close()

        if len(rows) == 0:
            flash("Email not found. Please SignUp", "error")
            return redirect(url_for('signup'))
        else:
             # Generate a random token
            reset_token = str(random.randint(100000, 999999))
            session['reset_token'] = reset_token
            subject = "Code For Password Change"
            body = f"Your verification code is {reset_token}. Please enter it on the website to verify your email."
            message = f"Subject: {subject}\n\n{body}"
            msg = send_mail(email, message)
            
            del reset_token,subject,body,message
            gc.collect()

            if msg == 0:
                flash("Code has been sent to registered email id.", "info")
                return redirect(url_for('validate_mail_reset_password', email=email))
            
        del email,sql,rows
        gc.collect()
    return render_template('forgot_password.html')

@app.route('/login', methods=['GET', 'POST'])
def login_page():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        
        sql = f"select * from public.register_table where email_id = '{email}';"
        curr,connection = sql_connection()
        curr.execute(sql)
        rows  = curr.fetchall()
        connection.close()

        if  len(rows) == 0 :
            flash("Email Id Not Found.", "error")

        if len(rows) != 0 :
            if rows[0][-1] != email :
                flash("Invalid Email Id", "error")
                return redirect(url_for('login_page'))
    
            decPassword = base64.b64decode(rows[0][-2]).decode("utf-8")
            if password == decPassword:
                session['email'] = email
                return redirect(url_for('chatpage'))
            else:
                flash("Invalid Password", "error")
            del [decPassword]
            gc.collect()
        del email,password,sql,rows
        gc.collect()
    return render_template('login.html')


@app.route('/chatpage')
def chatpage():
    
    try:
        if 'email' not in session:  # Check if user is logged in
            flash('Please login to access the chat page.', 'warning')  # Flash the message
            return redirect(url_for('login_page'))  # Redirect to login page

        # Allowed file extensions
        allowed_extensions = ('.xlsx', '.docx','.pdf','.txt','.jpg','.jpeg','.png','.csv')

        # credentials = authenticate_and_list_files()
        # Build Google Drive API client
        service = build('drive', 'v3', credentials=credentials)
        FOLDER_ID = '1SNEOT6spU3wD7AVJI_w53bfMDig2Ss4l' 
        # FOLDER_ID = '1LSgEkDLL8ulP7Ep-qvbqI2XbKQXf2oui11111'

        try:
            folder_metadata = service.files().get(fileId=FOLDER_ID, fields='id').execute()
        except HttpError as e:
            if e.resp.status == 404:
                raise FolderNotAvailable()
            else:
                raise e  # Re-raise other HttpError exceptions

        # Query files within the folder
        query = f"'{FOLDER_ID}' in parents"
        if service:
            files = []
            next_page_token = None

            # Paginate through the results
            while True:
                response = service.files().list(
                    q=query,
                    spaces='drive',
                    fields='nextPageToken, files(id, name)',
                    pageToken=next_page_token
                ).execute()

                files.extend(response.get('files', []))
                next_page_token = response.get('nextPageToken')

                if not next_page_token:
                    break
            # Print all files in the folder
            if files:
                files_with_id_dict = {}
                for file in files:
                    files_with_id_dict[file['name']] = file['id']
                    print(f"File Name: {file['name']}, File ID: {file['id']}")
                with open(os.path.join('Document','doc_files.json'), 'w') as fp:
                    json.dump(files_with_id_dict, fp)

                # Get the list of files with allowed extensions
                files = [os.path.basename(f) for f in [key for key,val in files_with_id_dict.items()]if f.endswith(allowed_extensions)]
                
                return render_template('chatpage.html', files=files)
            else:
                raise FolderNotAvailable()
    except FolderNotAvailable as exe:
        # flash('FolderNotAvailable','error')
        return jsonify({'error': str(exe)}), 404

@app.route('/uploads', methods=['POST'])
def upload_file():
    """Handle file uploads."""
    allowed_extensions = ('.csv', '.pdf', '.xlsx', '.txt')
    
    if 'file' not in request.files:
        flash('No file part','error')
        return redirect(url_for('index'))
    
    file = request.files['file']
    
    if file.filename == '':
        flash('No selected file','error')
        return redirect(url_for('index'))
    
    if file and file.filename.lower().endswith(allowed_extensions):
        filename = file.filename
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
        flash('File uploaded successfully','success')
    else:
        flash(f'Invalid file type. Only {", ".join(allowed_extensions)} files are allowed.','info')
    del allowed_extensions,file
    gc.collect()
    return redirect(url_for('index'))

@app.route('/delete/<filename>')
def delete_file(filename):
    """Delete a file.
    Args:
    filename(['String']) : Path of file to be deleted.
    """
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    os.remove(file_path)
    flash('File deleted successfully.','success')
    pkl_file_name = filename.split('.')[0]+'.pkl'
    if os.path.exists(os.path.join(pkl_file_name)):                
        os.remove(os.path.join(pkl_file_name))
    del file_path,pkl_file_name
    gc.collect()
    return redirect(url_for('index'))

 
# general prompt for model
prompt_temp = '''You are a Guide AI. Your role is to provide guidance strictly from information provided to you Dont go beyond information. Always respond in a concise and professional manner.

Follow these rules to handle user interactions:  
---
### **Rules for User Interactions**  
1. Must Do : 
   -include  URL or Link if provided in context as a reference for the user to find additional information if required else no need to show an url.

1a. **Gratitude or Acknowledgment (e.g., "Thank you," "Thanks," "Great answer"):**  
   - Respond only with:  
     *"I'll do my best to assist you. You're welcome! How can I assist you further?"*  
   - Do not include any additional context or information.  

2. **User Greets (e.g., "Hello," "Hi," "How are you?")**  
   - Respond politely and acknowledge. Examples:  
     - "Hello! How can I assist you today?"  
     - "Hi! How can I help you?"  

3. **Queries Based on Context:**  
   - Answer strictly based on the provided context.  
   - If the query is unrelated or the information is unavailable, respond with:  
     *"The information is not available in the provided context. Please provide more details or connect with a team member."*  

4. **Personal Information Shared by User:**  
   - Discourage sharing sensitive personal details. Example:  
     *"Thank you for sharing, but please avoid disclosing personal information."*  

5. **Full Forms or Definitions:**  
   - Provide full forms only if available in the context.  
   - If unavailable, respond with: *"Information not available in the provided context."*  

6. **Gratitude After Providing an Answer:**  
   - Respond with: *"You're welcome! Feel free to ask if you have any other questions."*  

7. **Unrelated or Unclear Questions:**  
   - If the query is unclear or unrelated, respond with:  
     *"The information is not available in the provided context. Please provide more details or connect with a team member."*  

8. **Response Formatting:**  
   - For gratitude, always use:  
     *"I'll do my best to assist you. You're welcome! How can I assist you further?"*  
   - For standard answers, use: `Answer: [Your Response]`.  
   - For unavailable information, use: `Answer: Not Found.`  
        example 
        1) any question related to document, but not provided to you then, respond with: *"Not Found"*
        2) Question like 'nbskbdjkdcdckdc ?' or 'nsxsxjnsakasniwq jsabxjsabjcbb wdjbsxqb ?', 'ueeygefbcbd akkahhidjn!'then, respond with: *"Not Found"*
9. Must Do : Whenever you cannot provide a complete answer to a question, include the provided URL as a reference for the user to find additional information if required.

---
### **Examples**  
1. **User Says "Thank You":**  
   - AI Reply: *"I'll do my best to assist you. You're welcome! How can I assist you further?"*  

2. **User Asks a Question with Greeting (e.g., "Hi, can you provide details about X?"):**  
   - AI Reply: *"Hi! Here's the information you requested about X..."*  

3. **User Asks for a Full Form (e.g., "What does DL stand for?"):**  
   - AI Reply: *"DL stands for Deep Learning."*  

4. **User Expresses Gratitude After an Answer:**  
   - AI Reply: *"You're welcome! Feel free to ask if you have any other questions."*  

5. **User Asks an Unrelated or Unclear Question:**  
   - AI Reply: *"The information is not available in the provided context. Please provide more details or connect with a team member."*  
---

**Context:**  
{context}  
{chat_history}  
Human: {question}  
'''

embedings = HuggingFaceBgeEmbeddings()
# model
# llm = ChatGroq(model='llama-3.1-70b-versatile',api_key='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',temperature=0,max_retries=2)

@app.route('/ask', methods=['POST','GET'])
def get_ans_from_csv():
    ''' this function is used to get answer from given csv.

    Args:
    doc_file([CSV]): comma separated file 
    query_text :  Question

    Returns: Answer
    '''
    if 'email' not in session:
        flash("Please log in to access this functionality.", "error")
        return redirect(url_for('login_page'))

    email = session['email']
    query_text = request.form.get('query_text')
    # doc_file = request.form.get('selected_file')
    selected_language = request.form.get('selected_language')
    # query_text = query_text.lower()
    if query_text :
        #to load pickle file 
        doc_file = 'anc.xlsx'
        pickle_file_name = doc_file.split('.')[0]+'.pkl'
        if os.path.isfile(os.path.join('pkl_files', pickle_file_name)):
            with open(os.path.join('pkl_files',pickle_file_name),mode='rb') as f:
                vector_index = pickle.load(f)
        else:
            # credentials = authenticate_and_list_files()
            # documents_id = '1hxi7y7uRS9gGwA9rBJE7ihLZJqGpLtS_'
            documents_id = '1Cak9V_QRjEbnlJzeEe8uj-VN864IrBpt'
            data = load_file_data(documents_id, credentials)

            if data == 'Not Found':
                res_ans = "Failed to load file from google drive.Please connect with Team Member"
                flash(res_ans, "error")
                return jsonify({'answer': res_ans,'question':query_text})

            else:
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=700,chunk_overlap=100)
                docs = text_splitter.split_documents(data)
                
                vector_index = FAISS.from_documents(docs, embedding=embedings)
                if not os.path.isfile(os.path.join('pkl_files',pickle_file_name)):
                    with open(os.path.join('pkl_files',pickle_file_name),mode='wb') as f:
                        pickle.dump(vector_index,f)
                
        # model
        llm = ChatGroq(model='llama-3.1-70b-versatile',api_key='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',temperature=0,max_retries=2)
        chat_memory = ConversationSummaryBufferMemory(llm=llm,memory_key='chat_history',return_messages=True)
        # chat_memory = ConversationSummaryBufferMemory(llm=llm,memory_key='chat_history',return_messages=True)
        prompt =  PromptTemplate(template=prompt_temp,input_variables=['context','chat_history','question'])
        # function is used to get answer
        chain = get_chain(llm,prompt,vector_index,chat_memory)
        res_dict = get_answer(chain,query_text,email,chat_memory)
        que_ans_dict = {'doc': doc_file, query_text:res_dict}
        res_ans =que_ans_dict[query_text][selected_language]

        del vector_index,chain,llm,res_dict,que_ans_dict,prompt
        gc.collect()
        
        return jsonify({'answer': res_ans,'question':query_text})
    else:
        return redirect(url_for('index'))

@app.route('/save_answers', methods=['POST'])
def save_answers():
    """Save answer data to answers.json."""
    data = request.json
    with open(answer_file_path, 'r+') as f:
        answers = json.load(f)
        answers.append(data)  # Append new data
        f.seek(0)  # Move to the beginning of the file
        json.dump(answers, f, indent=4)  # Save updated data
    del data
    gc.collect()
    return jsonify({'message': 'Answer data saved successfully'}), 200

@app.route('/save_feedback', methods=['POST'])
def save_feedback():
    """Save user feedback."""
    email= session['email']
    feedback_data = request.json

    sql_query = f"""select feedback from public.user_que_ans where email_id = '{email}';"""
    curr,conn = sql_connection()
    curr.execute(sql_query)
    res = curr.fetchall()
    conn.close()
    if res[0][0] is None:
        feed_back_lst = [feedback_data]
        feedback_json = json.dumps(feed_back_lst)
        sql_query = f"update public.user_que_ans set feedback = '{feedback_json}' where email_id = '{email}';"
        curr,conn = sql_connection()
        curr.execute(sql_query)
        conn.commit()
        conn.close()
    else:
        ans_lst = json.loads(res[0][0])
        ans_lst.append(feedback_data)
        feedback_json = json.dumps(ans_lst)
        # to excape the single quote , pgadmin gives error 
        feedback_json_escaped = feedback_json.replace("'", "''")

        sql_query = f"""update public.user_que_ans set feedback = '{feedback_json_escaped}'
                        where email_id = '{email}'
                    """
        curr,conn = sql_connection()
        curr.execute(sql_query)
        conn.commit()
        conn.close()

    return jsonify({'message': 'Feedback saved successfully'}), 200



@app.route('/clear', methods=['POST'])
def clear():
    """Clear user feedback or session data."""
    # Any session or data clearing logic goes here (if needed)

    # Redirect to chat page
    return redirect(url_for('chatpage'))


if __name__=='__main__':
    app.run(host='0.0.0.0',port=5011)







